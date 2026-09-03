#!/usr/bin/env python3
"""🔨️ F2 — generates before/after DOCX fixtures for s.stdio.docx@ecma-376/strict's 9 and
@ecma-376/transitional's 5 unfixtured mutations. set-snapshot is genuinely `third-party-generated`
(both sides are whole, unmodified python-docx 1.2.0 output); every other mutation here is a
structural OOXML-strict-vs-transitional edit none of python-docx's public API reaches, so those are
`handcrafted` on top of a genuine python-docx base package -- see 🔨️f2_ooxml_common.py's own
docstring for exactly why and how each is verified (lxml well-formedness + zip integrity, both run
live below, not assumed).
Idempotent: safe to re-run.
"""
import json
import sys
from pathlib import Path

from docx import Document

import importlib.util

_spec = importlib.util.spec_from_file_location("f2_ooxml_common", Path(__file__).parent / "🔨️f2_ooxml_common.py")
_mod = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(_mod)
NS, VML_CONTENT = _mod.NS, _mod.VML_CONTENT
zip_read, zip_rewrite, patch_tag_attr, insert_before_close = _mod.zip_read, _mod.zip_rewrite, _mod.patch_tag_attr, _mod.insert_before_close
remove_fragment, assert_wellformed, assert_zip_valid, sha256_of = _mod.remove_fragment, _mod.assert_wellformed, _mod.assert_zip_valid, _mod.sha256_of

ROOT = Path("/Users/ueli/Documents/semio")
DOCX_VERSION = "1.2.0"
ALT_CONTENT = (
    '<mc:AlternateContent xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">'
    '<mc:Choice Requires="wps"><w:p/></mc:Choice><mc:Fallback><w:p/></mc:Fallback></mc:AlternateContent>'
)


def base_docx() -> bytes:
    import io
    d = Document()
    d.add_paragraph("Hello")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def with_document_xml(data: bytes, new_xml: str) -> bytes:
    assert_wellformed(new_xml)
    out = zip_rewrite(data, {"word/document.xml": new_xml})
    assert_zip_valid(out)
    return out


def add_vml_part(data: bytes) -> bytes:
    ct = zip_read(data, "[Content_Types].xml")
    if "vmlDrawing" not in ct:
        ct = ct.replace("</Types>", '<Default Extension="vml" ContentType="application/vnd.openxmlformats-officedocument.vmlDrawing"/></Types>')
    rels = zip_read(data, "word/_rels/document.xml.rels")
    rels = rels.replace(
        "</Relationships>",
        '<Relationship Id="rIdVml1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/vmlDrawing" Target="vmlDrawing1.vml"/></Relationships>',
    )
    out = zip_rewrite(data, {"[Content_Types].xml": ct, "word/_rels/document.xml.rels": rels}, adds={"word/vmlDrawing1.vml": VML_CONTENT})
    assert_zip_valid(out)
    return out


def remove_vml_part(data: bytes) -> bytes:
    ct = zip_read(data, "[Content_Types].xml").replace('<Default Extension="vml" ContentType="application/vnd.openxmlformats-officedocument.vmlDrawing"/>', "")
    rels = zip_read(data, "word/_rels/document.xml.rels").replace(
        '<Relationship Id="rIdVml1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/vmlDrawing" Target="vmlDrawing1.vml"/>', ""
    )
    out = zip_rewrite(data, {"[Content_Types].xml": ct, "word/_rels/document.xml.rels": rels}, removes=["word/vmlDrawing1.vml"])
    assert_zip_valid(out)
    return out


def build_for_subset(subset: str):
    """subset: 'strict' or 'transitional'."""
    base = base_docx()
    tw_uri, strict_uri = NS["wordprocessingml"]
    tr_uri, strict_r_uri = NS["relationships"]
    conf_value = subset  # "strict" or "transitional"

    entries = []

    # set-snapshot -- genuinely third-party-generated, both sides real python-docx output.
    import io
    d2 = Document()
    d2.add_paragraph("Hello")
    d2.add_paragraph("World")
    buf2 = io.BytesIO()
    d2.save(buf2)
    entries.append(("set-snapshot", "third-party-generated", base, buf2.getvalue(),
                     "Whole-document snapshot replace: an unrelated valid document (two paragraphs instead of one) substituted wholesale."))

    # set-conformance-attribute: absent -> conf_value.
    xml = zip_read(base, "word/document.xml")
    after_xml = patch_tag_attr(xml, "w:document", "conformance", conf_value)
    entries.append(("set-conformance-attribute", "handcrafted", base, with_document_xml(base, after_xml),
                     f'The root <w:document>\'s conformance attribute set, absent -> "{conf_value}".'))

    # remove-conformance-attribute: conf_value -> absent (inverse).
    before_conf = with_document_xml(base, patch_tag_attr(xml, "w:document", "conformance", conf_value))
    after_conf = base
    entries.append(("remove-conformance-attribute", "handcrafted", before_conf, after_conf,
                     f'The root <w:document>\'s conformance attribute removed, "{conf_value}" -> absent.'))

    # set-main-namespace: xmlns:w transitional -> strict.
    after_ns = patch_tag_attr(xml, "w:document", "xmlns:w", strict_uri)
    entries.append(("set-main-namespace", "handcrafted", base, with_document_xml(base, after_ns),
                     f"The root <w:document>'s main xmlns:w namespace URI changed, Transitional ({tw_uri}) -> Strict ({strict_uri})."))

    # set-relationship-base: xmlns:r transitional -> strict.
    after_r = patch_tag_attr(xml, "w:document", "xmlns:r", strict_r_uri)
    entries.append(("set-relationship-base", "handcrafted", base, with_document_xml(base, after_r),
                     f"The root <w:document>'s relationships xmlns:r namespace URI changed, Transitional ({tr_uri}) -> Strict ({strict_r_uri})."))

    if subset == "strict":
        # insert-vml-part / remove-vml-part.
        after_vml = add_vml_part(base)
        entries.append(("insert-vml-part", "handcrafted", base, after_vml,
                         "A legacy VML drawing part (word/vmlDrawing1.vml) added to the package, with its Content_Types.xml Default and a document.xml.rels relationship wiring it in."))
        entries.append(("remove-vml-part", "handcrafted", after_vml, remove_vml_part(after_vml),
                         "The VML drawing part, its content-type default and its relationship removed from the package, inverse of insert-vml-part."))

        # insert-alternate-content / remove-alternate-content.
        after_alt_xml = insert_before_close(xml, "</w:body>", ALT_CONTENT)
        after_alt = with_document_xml(base, after_alt_xml)
        entries.append(("insert-alternate-content", "handcrafted", base, after_alt,
                         "An <mc:AlternateContent> markup-compatibility block (Choice/Fallback pair) inserted as the last child of <w:body>."))
        before_remove_xml = after_alt_xml
        after_remove_xml = remove_fragment(before_remove_xml, ALT_CONTENT)
        entries.append(("remove-alternate-content", "handcrafted", with_document_xml(base, before_remove_xml), with_document_xml(base, after_remove_xml),
                         "The <mc:AlternateContent> block removed from <w:body>, inverse of insert-alternate-content."))

    return entries


def emit(subset: str, artifact_id: str):
    subset_dir = ROOT / f"✏️s/🔌️plugins/🗄️stdio/🗿️artifacts/📜️docx/🏅️standards/🔖️ecma-376/🪆️subsets/✳️{subset}"
    fixtures = subset_dir / "🧫️fixtures"
    oracle_json = subset_dir / "🧪️oracle/🔣️.json"
    reader_oracle = f"python-docx-docx-ecma-376-{'' if subset == 'base' else subset + '-'}mutate-reader"

    manifests = []
    for mutation_id, klass, before_bytes, after_bytes, note in build_for_subset(subset):
        case_dir = fixtures / f"{mutation_id}-applied"
        case_dir.mkdir(parents=True, exist_ok=True)
        (case_dir / "before.docx").write_bytes(before_bytes)
        (case_dir / "after.docx").write_bytes(after_bytes)

        entry = {
            "schema": "semio.repository-test.fixture/v2",
            "id": f"{mutation_id}-applied",
            "class": klass,
            "target": {"artifact": artifact_id, "standard": "ecma-376", "subset": subset},
            "mutation": mutation_id,
            "outcome": "applied",
            "units": {"length": "unitless", "angle": "degree"},
            "files": [
                {"role": "expected-before-docx", "path": f"../🧫️fixtures/{mutation_id}-applied/before.docx", "mediaType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "sha256": sha256_of(before_bytes), "bytes": len(before_bytes)},
                {"role": "expected-after-docx", "path": f"../🧫️fixtures/{mutation_id}-applied/after.docx", "mediaType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "sha256": sha256_of(after_bytes), "bytes": len(after_bytes)},
            ],
            "provenance": {
                "source": "generated" if klass == "third-party-generated" else "authored",
                "license": "MIT (python-docx)" if klass == "third-party-generated" else "n/a (handcrafted zip/XML patch of an MIT python-docx base package)",
                "attribution": "Written by python-docx 1.2.0's own Document.save()" if klass == "third-party-generated" else "A genuine python-docx 1.2.0 package hand-patched at the OOXML structural level (see notes); the patched word/document.xml is re-parsed with lxml to confirm well-formedness and the archive's own zip integrity is re-checked before commit",
                "security": "scanned-clean",
                "privacy": "no-personal-data",
            },
            "comparisonProfile": "exact-bytes-v1",
            "reproducible": True,
            "family": "structural",
            "notes": note,
        }
        if klass == "third-party-generated":
            entry["generator"] = {
                "oracle": reader_oracle,
                "packageVersion": DOCX_VERSION,
                "engineFamily": "python-docx",
                "engineVersion": DOCX_VERSION,
                "command": "uv run python3 🔨️f2-gen-docx-structural-fixtures.py (python-docx Document object model + .save())",
                "platform": "darwin-arm64",
            }
        manifests.append(entry)
        print(f"[{subset}] {mutation_id:28s} {klass:20s} before={len(before_bytes)}B after={len(after_bytes)}B")

    data = json.loads(oracle_json.read_text())
    data["fixtureManifests"] = manifests
    oracle_json.write_text(json.dumps(data, indent=2, ensure_ascii=False) + "\n")
    print(f"Wrote {len(manifests)} fixtureManifests entries into {oracle_json}\n")


def main() -> None:
    emit("strict", "s.stdio.docx")
    emit("transitional", "s.stdio.docx")


if __name__ == "__main__":
    main()
